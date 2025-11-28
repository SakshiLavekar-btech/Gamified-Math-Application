from docx import Document
from docx.shared import Pt

# Create document
doc = Document()
doc.add_heading("MathWave: The Application For Creative Mindset", level=1)

# Cover Page
doc.add_paragraph("Mini Project Report Submitted by:\n\nSakshi Nitin Lavekar (2023BEC126)\n\nUnder the Guidance of:\nDr. Lenina Birgale")
doc.add_paragraph("\nDepartment of Electronics and Telecommunication Engineering\nSGGS Institute of Engineering and Technology, Vishnupuri, Nanded")

# Certification
doc.add_heading("CERTIFICATION", level=2)
doc.add_paragraph("""This is to certify that the project report titled “MathWave: The Application For Creative Wave” 
has been carried out by Sakshi Nitin Lavekar, a student of SGGS, in partial fulfillment of the requirements 
for the mini project of B.Tech Third Year under the guidance of Dr. Lenina Birgale. 
The project work is original and has not been submitted elsewhere.""")
doc.add_paragraph("\nGuided by:\nDr. Lenina Birgale")

# Declaration
doc.add_heading("DECLARATION", level=2)
doc.add_paragraph("""I hereby declare that the project report entitled “MathWave: The Application For Creative Wave” 
is a record of my original work under the supervision of Dr. Lenina Birgale. 
The contents of this report are based on my study. Any material obtained from external sources has been duly referenced.""")
doc.add_paragraph("\nSakshi Nitin Lavekar\n2023BEC126")

# Abstract
doc.add_heading("ABSTRACT", level=2)
doc.add_paragraph("""MathWave, as the name suggests, is an innovative solution for studying mathematics in a fun, 
engaging, and accessible way. This project reimagines learning as a process rooted in understanding and creative 
thinking—essential elements for the holistic development of every individual. The application offers a variety of 
interactive and playful methods to simplify complex mathematical concepts, making them easier to grasp and more 
enjoyable to explore. Through full-stack development, the project successfully integrates intuitive design, dynamic 
content, and responsive features to create a seamless learning experience. MathWave not only supports academic 
growth but also fosters curiosity, confidence, and a love for problem-solving among learners.""")

# Table of Contents
doc.add_heading("TABLE OF CONTENTS", level=2)
toc = [
"Chapter 1: Figures of Components Used",
"Chapter 2: Introduction & Problem Statement",
"Chapter 3: Construction and Implementation",
"Chapter 4: Methodology",
"Chapter 5: Block Diagram",
"Chapter 6: Diagrammatic Explanation of Working",
"Chapter 7: Conclusion and Future Scope",
"Chapter 8: References"
]
for item in toc:
    doc.add_paragraph(item, style="List Number")

# Introduction
doc.add_heading("Chapter 2: Introduction", level=2)
doc.add_paragraph("""We live in a world where technology is evolving at an unprecedented pace. From the early days when 
the internet was developed for military purposes to today—where it’s a household necessity—digital tools have become 
deeply embedded in our daily lives. However, this digital shift has also led to challenges, especially among children 
who spend excessive time on mobile devices, often at the cost of academic performance.""")
doc.add_paragraph("""The MathWave app was developed to transform this dynamic. It’s a fun and engaging educational game 
that turns screen time into learning time. With interactive quizzes, playful animations, and creative drawing features, 
the app makes studying feel like a game. Even users who typically avoid academic tasks find themselves drawn in—learning 
becomes irresistible. It’s not just an app; it’s a joyful blend of fun and education.""")

# Problem Statement
doc.add_heading("2.1 Problem Statement", level=3)
doc.add_paragraph("""Mathematics is often considered the most dreaded subject among young learners. Many students struggle 
because their foundational understanding is weak. MathWave was created to address this challenge by focusing on building 
core mathematical skills through engaging activities, interactive games, and visual learning tools.""")

# Construction
doc.add_heading("Chapter 3: Construction and Implementation", level=2)
doc.add_paragraph("""The MathWave web application was built using both frontend and backend technologies to ensure a 
smooth, interactive user experience. The frontend uses HTML, CSS, and JavaScript to create a responsive and dynamic 
interface. The backend, powered by Python (Flask/Django), manages user data, progress, and application logic.""")

# Methodology
doc.add_heading("Chapter 4: Methodology", level=2)
doc.add_paragraph("""MathWave was developed using a design thinking approach. User feedback from students and teachers 
was collected and analyzed to create an intuitive interface. Gamification principles were applied to improve engagement, 
and features like voice synthesis and animations were added for inclusivity and entertainment.""")

# Conclusion & Future Scope
doc.add_heading("Chapter 7: Conclusion and Future Scope", level=2)
doc.add_paragraph("""MathWave transforms learning from a passive activity into an engaging adventure. It demonstrates 
how technology, when applied thoughtfully, can make education both fun and effective. In the future, the app can be 
expanded to include multiplayer math challenges, AI-based tutoring, adaptive difficulty levels, and progress tracking 
to personalize learning experiences.""")

# References
doc.add_heading("Chapter 8: References", level=2)
doc.add_paragraph("""1. Arduino. Wikipedia: The Free Encyclopedia.\n2. Dhaief, Zahraa. (2016). People Counting Technology.\n3. Saxena, D. Songara. "Design of People Counting System," International Conference on Contemporary Computing (IC3), 2017.\n4. https://www.researchgate.net/figure/Link-to-the-profile-page_tbl2_221501358""")

# Save document
file_path = "/mnt/data/MathWave_Mini_Project_Report.docx"
doc.save(file_path)
file_path